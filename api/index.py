import os
import logging
import sys
from flask import Flask, jsonify, request, send_from_directory, render_template_string, redirect, url_for, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_admin import Admin, AdminIndexView, expose
from flask_admin.base import MenuLink
from flask_admin.contrib.sqla import ModelView
from flask_admin.form.upload import FileUploadField
from flask_cors import CORS
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from dotenv import load_dotenv
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure the local directory is in the path for Vercel
sys.path.append(os.path.abspath(os.path.dirname(__file__)))
from models import db, Profile, Project, Skill, SocialLink, Experience, Education, Certification

# Configuration
BASE_DIR = os.path.abspath(os.path.dirname(__file__))

# Load environment variables from api/.env
load_dotenv(os.path.join(BASE_DIR, '.env'))

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'a-very-secret-key-12345')

# Admin Credentials
# Fallbacks included for ease of use during initial setup
ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME", "saadkhi_uit").strip()
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "03353948753_gs150se").strip()

logger.info(f"DEBUG: ADMIN_USERNAME is {'SET (length: ' + str(len(ADMIN_USERNAME)) + ')' if ADMIN_USERNAME else 'NOT SET'}")
logger.info(f"DEBUG: ADMIN_PASSWORD is {'SET (length: ' + str(len(ADMIN_PASSWORD)) + ')' if ADMIN_PASSWORD else 'NOT SET'}")

# Handle Database Persistence
DATABASE_URL = os.environ.get('DATABASE_URL')

if DATABASE_URL:
    # Handle SQLAlchemy 1.4+ compatibility for 'postgres://' vs 'postgresql://'
    if DATABASE_URL.startswith("postgres://"):
        DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)
    
    # Ensure SSL for Neon/Postgres if not specified
    if "sslmode=" not in DATABASE_URL:
        separator = "&" if "?" in DATABASE_URL else "?"
        DATABASE_URL += f"{separator}sslmode=require"
        
    app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL
    logger.info("Using external persistent database (PostgreSQL/Neon)")
    UPLOAD_BASE_DIR = '/tmp/media' if os.environ.get('VERCEL') else os.path.join(BASE_DIR, 'media')
elif os.environ.get('VERCEL'):
    db_path = '/tmp/portfolio.db'
    repo_db_path = os.path.join(BASE_DIR, 'portfolio.db')
    if os.path.exists(repo_db_path) and not os.path.exists(db_path):
        import shutil
        try:
            shutil.copy2(repo_db_path, db_path)
            logger.info("Copied initial database to /tmp")
        except Exception as e:
            logger.error(f"Failed to copy database: {e}")
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
    UPLOAD_BASE_DIR = '/tmp/media'
else:
    app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{os.path.join(BASE_DIR, 'portfolio.db')}"
    UPLOAD_BASE_DIR = os.path.join(BASE_DIR, 'media')

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Ensure Media Directories Exist
for subpath in ['resumes', 'projects', 'skills', 'social_icons', 'profile']:
    path = os.path.join(UPLOAD_BASE_DIR, subpath)
    if not os.path.exists(path):
        try:
            os.makedirs(path, exist_ok=True)
            logger.info(f"Created directory: {path}")
        except Exception as e:
            logger.error(f"Failed to create directory {path}: {e}")
app.config['FLASK_ADMIN_SWATCH'] = 'cerulean'

# User model for session management
class User(UserMixin):
    def __init__(self, id):
        self.id = id

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    if user_id == ADMIN_USERNAME:
        return User(user_id)
    return None

# Custom Admin View with Security
class SecureModelView(ModelView):
    def is_accessible(self):
        return current_user.is_authenticated

    def _handle_view(self, name, **kwargs):
        if not self.is_accessible():
            return redirect(url_for('login', next=request.url))

class SecureAdminIndexView(AdminIndexView):
    def is_accessible(self):
        return current_user.is_authenticated

    def _handle_view(self, name, **kwargs):
        if not self.is_accessible():
            return redirect(url_for('login', next=request.url))

    @expose('/')
    def index(self):
        profiles = Profile.query.all()
        projects = Project.query.all()
        skills = Skill.query.all()
        experiences = Experience.query.all()
        educations = Education.query.all()
        certifications = Certification.query.all()
        is_ephemeral = os.environ.get('VERCEL') is not None and os.environ.get('DATABASE_URL') is None
        return self.render('admin/index.html', 
                           profiles=profiles, 
                           projects=projects, 
                           skills=skills, 
                           social_links=social_links,
                           experiences=experiences,
                           educations=educations,
                           certifications=certifications,
                           is_ephemeral=is_ephemeral)

# Initialize Database
db.init_app(app)

# Helper to ensure all tables and columns exist
def run_migrations():
    try:
        from sqlalchemy import inspect, text
        # Always run create_all to catch new tables
        db.create_all()
        
        inspector = inspect(db.engine)
        if 'profile' in inspector.get_table_names():
            columns = [c['name'] for c in inspector.get_columns('profile')]
            if 'profile_pic' not in columns:
                with db.engine.connect() as conn:
                    conn.execute(text('ALTER TABLE profile ADD COLUMN profile_pic VARCHAR(255)'))
                    conn.commit()
                    logger.info("Added profile_pic column to profile table")
            if 'resume_summary' not in columns:
                with db.engine.connect() as conn:
                    conn.execute(text('ALTER TABLE profile ADD COLUMN resume_summary TEXT'))
                    conn.commit()
                    logger.info("Added resume_summary column to profile table")
    except Exception as e:
        logger.error(f"Migration error: {e}")


# Database Seeding Logic
def seed_database():
    import json
    from models import Profile, Project, Skill, SocialLink
    
    # Only seed if Profile is empty
    if Profile.query.first():
        return
        
    data_file = os.path.join(BASE_DIR, 'data.json')
    if not os.path.exists(data_file):
        logger.warning(f"Seed data file {data_file} not found.")
        return

    try:
        with open(data_file, 'r') as f:
            data = json.load(f)
            
        logger.info("Seeding database from data.json...")
        
        # Profile
        p_data = data.get('profile', {})
        if p_data:
            profile = Profile(
                name=p_data.get('name'),
                title=p_data.get('title'),
                bio=p_data.get('bio'),
                resume_file=p_data.get('resume_file'),
                email=p_data.get('email'),
                location=p_data.get('location'),
                phone_number=p_data.get('phone_number')
            )
            db.session.add(profile)
            
        # Projects
        for proj in data.get('projects', []):
            p = Project(
                title=proj.get('title'),
                description=proj.get('description'),
                tech_stack=proj.get('tech_stack'),
                image=proj.get('image'),
                video=proj.get('video'),
                live_link=proj.get('live_link'),
                category=proj.get('category'),
                is_featured=bool(proj.get('is_featured'))
            )
            db.session.add(p)
            
        # Skills
        for skill in data.get('skills', []):
            s = Skill(
                name=skill.get('name'),
                icon=skill.get('icon'),
                order=skill.get('order', 0)
            )
            db.session.add(s)
            
        # Social Links
        for link in data.get('social_links', []):
            sl = SocialLink(
                name=link.get('name'),
                url=link.get('url'),
                icon_class=link.get('icon_class'),
                icon_image=link.get('icon_image'),
                order=link.get('order', 0)
            )
            db.session.add(sl)
            
        db.session.commit()
        logger.info("Database seeding completed.")
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error seeding database: {e}")

# Initialize Tables and Seed safely
if not os.environ.get('VERCEL'):
    with app.app_context():
        run_migrations()
        seed_database()
else:
    # On Vercel, apply a lazy initialization inside before_request to prevent halting the lambda container creation
    db_initialized = False
    @app.before_request
    def lazy_db_init():
        global db_initialized
        if not db_initialized:
            run_migrations()
            # Still check for seed if profile is missing
            try:
                if not Profile.query.first():
                    seed_database()
            except Exception:
                pass
            db_initialized = True

# Helper for thumbnail formatting
from markupsafe import Markup

def thumbnail_formatter(view, context, model, name):
    if not getattr(model, name):
        return ""
    
    # Standardize the path for display
    base_name = os.path.basename(getattr(model, name))
    # Determine subfolder based on model type
    subfolder = ""
    if isinstance(model, Profile): subfolder = "profile"
    elif isinstance(model, Project): subfolder = "projects"
    elif isinstance(model, Skill): subfolder = "skills"
    elif isinstance(model, SocialLink): subfolder = "social_icons"
    
    url = f"/media/{subfolder}/{base_name}"
    return Markup(f'<img src="{url}" height="40" style="border-radius: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">')

# Specialized ModelViews with FileUploadField/ImageUploadField
from flask_admin.form.upload import ImageUploadField

class ProfileModelView(SecureModelView):
    column_list = ('name', 'title', 'email', 'profile_pic_thumb')
    column_labels = {'profile_pic_thumb': 'Photo'}
    column_formatters = {
        'profile_pic_thumb': lambda v, c, m, p: thumbnail_formatter(v, c, m, 'profile_pic')
    }
    
    form_overrides = {
        'resume_file': FileUploadField,
        'profile_pic': ImageUploadField
    }
    form_args = {
        'resume_file': {
            'label': 'Resume (PDF)',
            'base_path': os.path.join(UPLOAD_BASE_DIR, 'resumes'),
            'allow_overwrite': True
        },
        'profile_pic': {
            'label': 'Profile Picture',
            'base_path': os.path.join(UPLOAD_BASE_DIR, 'profile'),
            'allow_overwrite': True
        }
    }

class ProjectModelView(SecureModelView):
    column_list = ('title', 'category', 'is_featured', 'image_thumb')
    column_labels = {'image_thumb': 'Preview'}
    column_formatters = {
        'image_thumb': lambda v, c, m, p: thumbnail_formatter(v, c, m, 'image')
    }
    
    form_overrides = {
        'image': ImageUploadField,
        'video': FileUploadField
    }
    form_args = {
        'image': {
            'label': 'Project Image',
            'base_path': os.path.join(UPLOAD_BASE_DIR, 'projects'),
            'allow_overwrite': True
        },
        'video': {
            'label': 'Project Video',
            'base_path': os.path.join(UPLOAD_BASE_DIR, 'projects'),
            'allow_overwrite': True
        }
    }

class SkillModelView(SecureModelView):
    column_list = ('name', 'order', 'icon_thumb')
    column_labels = {'icon_thumb': 'Icon'}
    column_formatters = {
        'icon_thumb': lambda v, c, m, p: thumbnail_formatter(v, c, m, 'icon')
    }
    
    form_overrides = {
        'icon': ImageUploadField
    }
    form_args = {
        'icon': {
            'label': 'Skill Icon',
            'base_path': os.path.join(UPLOAD_BASE_DIR, 'skills'),
            'allow_overwrite': True
        }
    }

class SocialLinkModelView(SecureModelView):
    column_list = ('name', 'url', 'icon_thumb')
    column_labels = {'icon_thumb': 'Icon'}
    column_formatters = {
        'icon_thumb': lambda v, c, m, p: thumbnail_formatter(v, c, m, 'icon_image')
    }
    
    form_overrides = {
        'icon_image': ImageUploadField
    }
    form_args = {
        'icon_image': {
            'label': 'Icon Image',
            'base_path': os.path.join(UPLOAD_BASE_DIR, 'social_icons'),
            'allow_overwrite': True
        }
    }

# Initialize Admin with Security
admin = Admin(app, name='Portfolio Admin', template_mode='bootstrap3', index_view=SecureAdminIndexView(), base_template='admin/custom_base.html')
admin.add_view(ProfileModelView(Profile, db.session))
admin.add_view(ProjectModelView(Project, db.session))
admin.add_view(SkillModelView(Skill, db.session))
admin.add_view(SocialLinkModelView(SocialLink, db.session))
admin.add_view(SecureModelView(Experience, db.session))
admin.add_view(SecureModelView(Education, db.session))
admin.add_view(SecureModelView(Certification, db.session))

# logout Link
admin.add_link(MenuLink(name='Logout', category='', url='/logout'))

# Authentication Routes
@app.route('/login', methods=['GET', 'POST'])
def login():
    error_msg = None
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            user = User(username)
            login_user(user)
            return redirect(request.args.get('next') or url_for('admin.index'))
        
        logger.warning(f"Login failed for user '{username}'.")
        error_msg = "Invalid credentials"
        if not ADMIN_USERNAME or not ADMIN_PASSWORD:
            error_msg = "Server Error: Admin credentials not configured"

    return render_template_string('''
        <!DOCTYPE html>
        <html>
        <head>
            <title>Admin Login | Portfolio</title>
            <meta name="viewport" content="width=device-width, initial-scale=1">
            <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600&family=Space+Grotesk:wght@600&display=swap" rel="stylesheet">
            <style>
                :root {
                    --primary: #ff8a00;
                    --primary-hover: #ffa733;
                    --bg: #f8fafc;
                    --card-bg: #ffffff;
                    --text: #1e293b;
                    --text-muted: #64748b;
                    --border: #e2e8f0;
                }
                body {
                    font-family: 'Inter', sans-serif;
                    background-color: var(--bg);
                    color: var(--text);
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    min-height: 100vh;
                    margin: 0;
                }
                .login-card {
                    background: var(--card-bg);
                    border: 1px solid var(--border);
                    padding: 3rem;
                    border-radius: 2rem;
                    width: 100%;
                    max-width: 420px;
                    box-shadow: 0 20px 25px -5px rgba(0, 0, 0, 0.05), 0 10px 10px -5px rgba(0, 0, 0, 0.02);
                }
                .header {
                    margin-bottom: 2.5rem;
                    text-align: center;
                }
                .header h1 {
                    font-size: 2.25rem;
                    font-weight: 700;
                    font-family: 'Space Grotesk', sans-serif;
                    margin: 0 0 0.5rem 0;
                    color: var(--primary);
                    letter-spacing: -1px;
                }
                .header p {
                    color: var(--text-muted);
                    font-size: 0.95rem;
                    margin: 0;
                }
                .form-group {
                    margin-bottom: 1.5rem;
                }
                .form-group label {
                    display: block;
                    font-size: 0.875rem;
                    font-weight: 500;
                    margin-bottom: 0.6rem;
                    color: var(--text-muted);
                    text-transform: uppercase;
                    letter-spacing: 0.05em;
                }
                input {
                    width: 100%;
                    padding: 1rem 1.25rem;
                    background: #f8fafc;
                    border: 1px solid var(--border);
                    border-radius: 1rem;
                    color: var(--text);
                    font-size: 1rem;
                    transition: all 0.3s;
                    box-sizing: border-box;
                }
                input:focus {
                    outline: none;
                    border-color: var(--primary);
                    background: #ffffff;
                    box-shadow: 0 0 0 4px rgba(255, 138, 0, 0.1);
                }
                button {
                    width: 100%;
                    padding: 1rem;
                    background: var(--primary);
                    color: white;
                    border: none;
                    border-radius: 1rem;
                    font-size: 1.1rem;
                    font-weight: 700;
                    cursor: pointer;
                    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                    margin-top: 1.5rem;
                }
                button:hover {
                    background: #ffa733;
                    transform: translateY(-2px);
                    box-shadow: 0 10px 20px rgba(255, 138, 0, 0.2);
                }
                button:active {
                    transform: translateY(0);
                }
                .error {
                    background: #fef2f2;
                    border: 1px solid #fee2e2;
                    color: #ef4444;
                    padding: 1rem;
                    border-radius: 1rem;
                    font-size: 0.9rem;
                    margin-bottom: 2rem;
                    text-align: center;
                }
            </style>
        </head>
        <body>
            <div class="login-card">
                <div class="header">
                    <h1>Admin Access</h1>
                    <p>Portfolio Dashboard Login</p>
                </div>
                
                {% if error %}
                <div class="error">
                    {{ error }}
                </div>
                {% endif %}
                
                <form method="post">
                    <div class="form-group">
                        <label>Username</label>
                        <input type="text" name="username" placeholder="admin" required autofocus>
                    </div>
                    <div class="form-group">
                        <label>Password</label>
                        <input type="password" name="password" placeholder="••••••••" required>
                    </div>
                    <button type="submit">Log In</button>
                </form>
            </div>
        </body>
        </html>
    ''', error=error_msg)

@app.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('login'))

# Serve Media Files
@app.route('/media/<path:filename>')
def serve_media(filename):
    # Search in multiple potential locations for compatibility
    search_dirs = [
        os.path.join(BASE_DIR, 'media'),
        '/tmp/media'
    ]
    
    # Files might be nested in subfolders like 'skills/icons/' due to legacy data
    # We strip any path parts and search for the basename in our standard subfolders
    base_name = os.path.basename(filename)
    
    # Standard subfolders to check
    subfolders = ['', 'resumes', 'projects', 'skills', 'social_icons', 'profile', 'projects/images', 'skills/icons']
    
    for d in search_dirs:
        if not os.path.exists(d):
            continue
        # 1. First check if the file exists directly or in specific subfolder
        for sub in subfolders:
            target_path = os.path.join(d, sub, base_name)
            if os.path.isfile(target_path):
                return send_from_directory(os.path.join(d, sub), base_name)
        
        # 2. Recursive search for robustness
        import glob
        matches = glob.glob(os.path.join(d, '**', base_name), recursive=True)
        if matches:
            return send_from_directory(os.path.dirname(matches[0]), base_name)

    return jsonify({"error": "Media file not found", "attempted": filename}), 404

@app.route('/api/portfolio/', methods=['GET'])
@app.route('/api/portfolio', methods=['GET'])
def get_portfolio():
    try:
        profile = Profile.query.first()
    except Exception as e:
        logger.error(f"Error fetching profile: {e}")
        profile = None
    
    # Use database values or defaults for profile info
    name = profile.name if profile else "Saad Ali"
    title = profile.title if profile else "AI & Software Engineer"
    bio = profile.bio if profile else "I'm a software engineer focused on building clean, scalable backends..."
    email = profile.email if profile else "saadalioffic@gmail.com"
    phone = profile.phone_number if profile else ""
    location = profile.location if profile else "Karachi, Pakistan"
    
    resume_url = ""
    if profile and profile.resume_file:
        resume_url = f"/media/resumes/{os.path.basename(profile.resume_file)}" if not str(profile.resume_file).startswith('/') else profile.resume_file
        
    profile_pic = ""
    if profile and profile.profile_pic:
        profile_pic = f"/media/profile/{os.path.basename(profile.profile_pic)}" if not str(profile.profile_pic).startswith('/') else profile.profile_pic

    projects = Project.query.order_by(Project.created_at.desc()).all()
    skills = Skill.query.order_by(Skill.order).all()
    social_links = SocialLink.query.order_by(SocialLink.order).all()

    featured_projects = []
    latest_projects = []
    
    for p in projects:
        p_dict = {
            "id": p.id,
            "title": p.title,
            "description": p.description,
            "tech_stack": p.tech_stack or "",
            "image": f"/media/projects/{os.path.basename(p.image)}" if p.image and not p.image.startswith('http') and not p.image.startswith('/') else p.image,
            "video": f"/media/projects/{os.path.basename(p.video)}" if p.video and not p.video.startswith('http') and not p.video.startswith('/') else p.video,
            "live_link": p.live_link or "#",
            "category": p.category or "Development",
            "is_featured": p.is_featured
        }
        if p.is_featured:
            featured_projects.append(p_dict)
        else:
            latest_projects.append(p_dict)

    response = jsonify({
        "hero": {
            "name": name,
            "title": title,
            "subtitle": "AI & Software Engineer",
            "cta_primary": "See My Work",
            "cta_secondary": "Get in Touch",
            "resume_url": resume_url,
            "profile_pic": profile_pic
        },
        "about": {
            "title": "About Me",
            "description": bio
        },
        "skills": [{
            "id": s.id,
            "name": s.name,
            "icon": f"/media/skills/{os.path.basename(s.icon)}" if s.icon and not s.icon.startswith('http') and not s.icon.startswith('fa') and not s.icon.startswith('/') else (s.icon or ""),
            "order": s.order
        } for s in skills],
        "social_links": [{
            "id": sl.id,
            "name": sl.name,
            "url": sl.url,
            "icon_class": sl.icon_class or "fa-solid fa-link",
            "icon_image": f"/media/social_icons/{os.path.basename(sl.icon_image)}" if sl.icon_image and not sl.icon_image.startswith('/') else (sl.icon_image or "")
        } for sl in social_links],
        "featured_projects": featured_projects,
        "latest_projects": latest_projects,
        "contact": {
            "email": email,
            "phone": phone,
            "location": location,
            "linkedin": "https://linkedin.com/in/saadkhi",
            "github": "https://github.com/saadkhi"
        }
    })
    
    # Enable edge caching for fast loading in production (1 hour max-age, 1 day stale-while-revalidate)
    response.headers["Cache-Control"] = "public, s-maxage=3600, stale-while-revalidate=86400"
    
    return response

@app.route('/api/projects/', methods=['GET'])
@app.route('/api/projects', methods=['GET'])
def get_projects():
    projects = Project.query.order_by(Project.created_at.desc()).all()
    response = jsonify([{
        "id": p.id,
        "title": p.title,
        "description": p.description,
        "tech_stack": p.tech_stack,
        "image": f"/media/projects/{os.path.basename(p.image)}" if p.image and not p.image.startswith('http') and not p.image.startswith('/') else p.image,
        "video": f"/media/projects/{os.path.basename(p.video)}" if p.video and not p.video.startswith('http') and not p.video.startswith('/') else p.video,
        "live_link": p.live_link,
        "category": p.category,
        "is_featured": p.is_featured
    } for p in projects])
    
    # Enable edge caching
    response.headers["Cache-Control"] = "public, s-maxage=3600, stale-while-revalidate=86400"
    
    return response

@app.route('/api/resume/', methods=['GET'])
@app.route('/api/resume', methods=['GET'])
def get_resume():
    try:
        profile = Profile.query.first()
    except Exception as e:
        logger.error(f"Error fetching profile for resume: {e}")
        profile = None

    try:
        experiences = Experience.query.order_by(Experience.order).all()
    except Exception as e:
        logger.error(f"Error fetching experiences: {e}")
        experiences = []

    try:
        educations = Education.query.order_by(Education.order).all()
    except Exception as e:
        logger.error(f"Error fetching education: {e}")
        educations = []

    try:
        certifications = Certification.query.order_by(Certification.order).all()
    except Exception as e:
        logger.error(f"Error fetching certifications: {e}")
        certifications = []

    try:
        skills = Skill.query.order_by(Skill.order).all()
    except Exception as e:
        logger.error(f"Error fetching skills: {e}")
        skills = []

    try:
        projects = Project.query.order_by(Project.created_at.desc()).all()
    except Exception as e:
        logger.error(f"Error fetching projects: {e}")
        projects = []

    try:
        social_links = SocialLink.query.order_by(SocialLink.order).all()
    except Exception as e:
        logger.error(f"Error fetching social links: {e}")
        social_links = []

    resume_data = {
        "header": {
            "name": profile.name if profile else "Saad Ali",
            "title": profile.title if profile else "AI & Software Engineer",
            "email": profile.email if profile else "saadalioffic@gmail.com",
            "phone": profile.phone_number if profile else "",
            "location": profile.location if profile else "Karachi, Pakistan",
            "github": "https://github.com/saadkhi",
            "linkedin": "https://linkedin.com/in/saadkhi",
            "portfolio": request.host_url
        },
        "summary": profile.resume_summary if profile else "",
        "experiences": [{
            "id": e.id,
            "title": e.title,
            "company": e.company,
            "location": e.location,
            "period": e.period,
            "description": e.description
        } for e in experiences],
        "education": [{
            "id": ed.id,
            "degree": ed.degree,
            "university": ed.university,
            "year": ed.year,
            "gpa": ed.gpa,
            "courses": ed.courses
        } for ed in educations],
        "skills": [s.name for s in skills],
        "projects": [{
            "title": p.title,
            "description": p.description,
            "tech_stack": p.tech_stack,
            "live_link": p.live_link
        } for p in projects if p.is_featured],
        "certifications": [{
            "name": c.name,
            "issuer": c.issuer,
            "year": c.year
        } for c in certifications],
        "social_links": [{
            "name": sl.name,
            "url": sl.url
        } for sl in social_links]
    }
    
    return jsonify(resume_data)

@app.route('/api/resume/download', methods=['GET'])
def download_resume():
    profile = Profile.query.first()
    experiences = Experience.query.order_by(Experience.order).all()
    educations = Education.query.order_by(Education.order).all()
    certifications = Certification.query.order_by(Certification.order).all()
    skills = Skill.query.order_by(Skill.order).all()
    projects = Project.query.order_by(Project.created_at.desc()).all()

    doc = Document()
    
    # 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Header
    name = profile.name if profile else "Saad Ali"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name.upper())
    run.bold = True
    run.font.size = Pt(14) # Name slightly larger than headings
    run.font.name = 'Arial'

    contact_info = []
    if profile.location: contact_info.append(profile.location)
    if profile.phone_number: contact_info.append(profile.phone_number)
    if profile.email: contact_info.append(profile.email)
    
    p = doc.add_paragraph(' | '.join(contact_info))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"GitHub: {profile.name.lower().replace(' ', '')} | LinkedIn: {profile.name.lower().replace(' ', '')}")

    def add_section_heading(text):
        h = doc.add_paragraph()
        run = h.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'

    # Summary
    if profile.resume_summary:
        add_section_heading("Professional Summary")
        doc.add_paragraph(profile.resume_summary)

    # Skills
    if skills:
        add_section_heading("Technical Skills")
        doc.add_paragraph(', '.join([s.name for s in skills]))

    # Experience
    if experiences:
        add_section_heading("Work Experience")
        for exp in experiences:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.title} - {exp.company}")
            run.bold = True
            p.add_run(f"\t{exp.period}").italic = True
            
            p = doc.add_paragraph(exp.location)
            p.style.font.size = Pt(9)
            
            # Bullets
            for bullet in exp.description.split('\n'):
                if bullet.strip():
                    doc.add_paragraph(bullet.strip().replace('- ', '').replace('• ', ''), style='List Bullet')

    # Education
    if educations:
        add_section_heading("Education")
        for edu in educations:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.degree} - {edu.university}")
            run.bold = True
            p.add_run(f"\t{edu.year}").italic = True
            if edu.gpa or edu.courses:
                info = []
                if edu.gpa: info.append(f"GPA: {edu.gpa}")
                if edu.courses: info.append(f"Courses: {edu.courses}")
                doc.add_paragraph('. '.join(info))

    # Projects
    featured_projects = [p for p in projects if p.is_featured]
    if featured_projects:
        add_section_heading("Projects")
        for proj in featured_projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.title)
            run.bold = True
            doc.add_paragraph(proj.description)
            doc.add_paragraph(f"Tech: {proj.tech_stack}").italic = True

    # Certifications
    if certifications:
        add_section_heading("Certifications")
        for cert in certifications:
            doc.add_paragraph(f"{cert.name} ({cert.issuer}) - {cert.year}")

    # Save to buffer
    target = BytesIO()
    doc.save(target)
    target.seek(0)
    
    filename = f"{name.replace(' ', '_')}_Resume.docx"
    return send_file(target, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/contact/', methods=['POST'])
@app.route('/api/contact', methods=['POST'])
def contact_form_submission():
    data = request.json
    google_script_url = os.environ.get('GOOGLE_SHEETS_SCRIPT_URL')
    
    if not google_script_url:
        logger.error("GOOGLE_SHEETS_SCRIPT_URL not configured.")
        return jsonify({'error': True, 'message': 'Service temporarily unavailable.'}), 500
    
    try:
        response = requests.post(google_script_url, json=data, timeout=10)
        return jsonify({
            'error': False,
            'message': 'Message sent successfully!'
        })
    except requests.exceptions.RequestException as e:
        logger.error(f"Error proxying to Google Sheets: {str(e)}")
        return jsonify({'error': True, 'message': 'Failed to forward message.'}), 503

@app.route('/health')
def health_check():
    return jsonify({"status": "ok"})

# Removed redundant db.create_all() as it's handled above

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8000))
    app.run(host='0.0.0.0', port=port, debug=True)
